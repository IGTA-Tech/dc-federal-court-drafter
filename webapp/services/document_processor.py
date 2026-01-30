"""
Document Processor Service

Extracts content from uploaded DOCX files and reformats them
to DC Federal District Court standards.
"""
import re
from io import BytesIO
from datetime import datetime
from typing import Optional, Dict, List, Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))
from config import FORMAT_SPECS, DC_COURT_NAME, DOCUMENT_TYPES


# Keywords for auto-detecting sections
SECTION_KEYWORDS = {
    "introduction": [
        "introduction", "preliminary statement", "nature of the case",
        "comes now", "respectfully submits", "respectfully moves"
    ],
    "facts": [
        "factual background", "statement of facts", "facts",
        "background", "statement of the case", "procedural history"
    ],
    "legal_standard": [
        "legal standard", "standard of review", "applicable law",
        "governing law", "legal framework"
    ],
    "argument": [
        "argument", "discussion", "analysis", "reasons",
        "points and authorities"
    ],
    "conclusion": [
        "conclusion", "wherefore", "prayer for relief", "relief requested",
        "for the foregoing reasons"
    ]
}


class DocumentProcessor:
    """Processes uploaded documents and reformats them to DC court standards."""

    def extract_from_docx(self, file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text and structure from a DOCX file.

        Args:
            file_bytes: Raw bytes of the DOCX file

        Returns:
            Dictionary containing extracted content and metadata
        """
        doc = Document(BytesIO(file_bytes))

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect if paragraph is a heading
                is_heading = self._is_heading(para)
                paragraphs.append({
                    "text": text,
                    "is_heading": is_heading,
                    "style": para.style.name if para.style else "Normal"
                })

        # Try to extract case information from caption
        case_info = self._extract_case_info(paragraphs)

        # Auto-detect sections
        sections = self.auto_detect_sections(paragraphs)

        return {
            "paragraphs": paragraphs,
            "full_text": "\n\n".join([p["text"] for p in paragraphs]),
            "case_info": case_info,
            "sections": sections,
            "word_count": sum(len(p["text"].split()) for p in paragraphs),
            "paragraph_count": len(paragraphs)
        }

    def _is_heading(self, para) -> bool:
        """Determine if a paragraph is a heading."""
        # Check style name
        if para.style and "heading" in para.style.name.lower():
            return True

        # Check if all caps and short
        text = para.text.strip()
        if text.isupper() and len(text) < 100:
            return True

        # Check if bold throughout
        if para.runs:
            all_bold = all(run.bold for run in para.runs if run.text.strip())
            if all_bold and len(text) < 100:
                return True

        return False

    def _extract_case_info(self, paragraphs: List[Dict]) -> Dict[str, str]:
        """Extract case information from document text."""
        full_text = "\n".join([p["text"] for p in paragraphs[:20]])  # Check first 20 paragraphs

        case_info = {
            "case_number": "",
            "plaintiff": "",
            "defendant": "",
            "judge_name": "",
            "document_title": ""
        }

        # Extract case number (various formats)
        case_patterns = [
            r'Case\s*(?:No\.?|Number)[:\s]*(\d+[:\-]\d+\-cv\-\d+(?:\-[A-Z]+)?)',
            r'(\d+[:\-]\d+\-cv\-\d+(?:\-[A-Z]+)?)',
            r'Civil\s*(?:Action|Case)\s*(?:No\.?)?[:\s]*(\d+[\-:]\d+)',
        ]

        for pattern in case_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                case_info["case_number"] = match.group(1)
                break

        # Extract parties from "v." pattern
        vs_patterns = [
            r'([A-Z][A-Za-z\s,\.]+?)\s*,?\s*(?:Plaintiff|Petitioner).*?v\.?\s*([A-Z][A-Za-z\s,\.]+?)\s*,?\s*(?:Defendant|Respondent)',
            r'([A-Z][A-Z\s,\.]+)\s+v\.?\s+([A-Z][A-Z\s,\.]+)',
        ]

        for pattern in vs_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE | re.DOTALL)
            if match:
                case_info["plaintiff"] = match.group(1).strip()
                case_info["defendant"] = match.group(2).strip()
                break

        # Extract judge name
        judge_patterns = [
            r'(?:Judge|Hon\.?|Honorable)[:\s]+([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)',
            r'before\s+(?:the\s+)?(?:Hon\.?|Honorable)\s+([A-Z][a-z]+\s+[A-Z][a-z]+)',
        ]

        for pattern in judge_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                case_info["judge_name"] = match.group(1).strip()
                break

        # Try to find document title
        for para in paragraphs[:15]:
            if para["is_heading"]:
                text = para["text"].strip()
                title_keywords = ["motion", "opposition", "reply", "memorandum", "complaint", "answer", "brief"]
                if any(kw in text.lower() for kw in title_keywords):
                    case_info["document_title"] = text
                    break

        return case_info

    def auto_detect_sections(self, paragraphs: List[Dict]) -> Dict[str, str]:
        """
        Automatically detect and extract document sections.

        Args:
            paragraphs: List of paragraph dictionaries

        Returns:
            Dictionary with section names as keys and content as values
        """
        sections = {
            "introduction": "",
            "facts": "",
            "legal_standard": "",
            "argument": "",
            "conclusion": ""
        }

        current_section = None
        current_content = []

        for para in paragraphs:
            text = para["text"]
            text_lower = text.lower()

            # Check if this paragraph starts a new section
            detected_section = None
            if para["is_heading"]:
                for section_name, keywords in SECTION_KEYWORDS.items():
                    if any(kw in text_lower for kw in keywords):
                        detected_section = section_name
                        break

            if detected_section:
                # Save previous section content
                if current_section and current_content:
                    sections[current_section] = "\n\n".join(current_content)

                current_section = detected_section
                current_content = []
            elif current_section:
                # Add to current section
                current_content.append(text)
            else:
                # Before first section detected - treat as introduction
                if not sections["introduction"]:
                    # Skip caption-like content
                    if not self._looks_like_caption(text):
                        if "introduction" not in sections or not sections["introduction"]:
                            current_section = "introduction"
                            current_content.append(text)

        # Save last section
        if current_section and current_content:
            sections[current_section] = "\n\n".join(current_content)

        return sections

    def _looks_like_caption(self, text: str) -> bool:
        """Check if text looks like a caption block."""
        caption_indicators = [
            "UNITED STATES DISTRICT COURT",
            "DISTRICT OF COLUMBIA",
            "Plaintiff,",
            "Defendant,",
            "Case No.",
            "Civil Action",
            "_____"
        ]
        return any(ind in text for ind in caption_indicators)

    def reformat_to_dc_standards(
        self,
        content: Dict[str, Any],
        doc_type: str,
        case_info: Dict[str, str],
        attorney_info: Optional[Dict[str, str]] = None,
        include_certificate: bool = True
    ) -> bytes:
        """
        Reformat extracted content to DC District Court standards.

        Args:
            content: Extracted content from document
            doc_type: Type of document (motion_to_dismiss, opposition, etc.)
            case_info: Case information (case_number, plaintiff, defendant, judge)
            attorney_info: Attorney information for signature block
            include_certificate: Whether to include certificate of service

        Returns:
            Bytes of the reformatted DOCX file
        """
        doc = Document()

        # Set up document formatting per LCvR 7(o)(1)
        self._setup_document_format(doc)

        # Add caption
        self._add_caption(doc, case_info)

        # Add document title
        self._add_document_title(doc, doc_type, case_info.get("custom_title", ""))

        # Add body content with proper formatting
        sections = content.get("sections", {})
        self._add_body_content(doc, sections, case_info)

        # Add signature block
        if attorney_info:
            self._add_signature_block(doc, attorney_info, case_info)

        # Add certificate of service
        if include_certificate and attorney_info:
            self._add_certificate_of_service(doc, attorney_info, case_info)

        # Add page numbers
        self._add_page_numbers(doc)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def _setup_document_format(self, doc: Document):
        """Set up document formatting per LCvR 7(o)(1)."""
        section = doc.sections[0]

        # Page size: Letter
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        # Margins: 1 inch all sides
        section.top_margin = Inches(FORMAT_SPECS["margin_inches"])
        section.bottom_margin = Inches(FORMAT_SPECS["margin_inches"])
        section.left_margin = Inches(FORMAT_SPECS["margin_inches"])
        section.right_margin = Inches(FORMAT_SPECS["margin_inches"])

        # Default font: Times New Roman 12pt
        style = doc.styles['Normal']
        font = style.font
        font.name = FORMAT_SPECS["font_name"]
        font.size = Pt(FORMAT_SPECS["font_size"])

        # Line spacing: Double
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    def _add_caption(self, doc: Document, case_info: Dict[str, str]):
        """Add the caption block per LCvR 5.1(b)(c)."""
        # Court name - centered, bold
        court_para = doc.add_paragraph()
        court_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        court_run = court_para.add_run(DC_COURT_NAME)
        court_run.bold = True
        court_run.font.name = FORMAT_SPECS["font_name"]
        court_run.font.size = Pt(FORMAT_SPECS["font_size"])

        doc.add_paragraph()

        # Create caption table
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Inches(3.5)
        table.columns[1].width = Inches(2.5)

        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]

        # Build parties block
        plaintiff = case_info.get("plaintiff", "PLAINTIFF NAME").upper()
        defendant = case_info.get("defendant", "DEFENDANT NAME").upper()

        parties_lines = [
            "________________________________________",
            "                                        )",
            f"{plaintiff},",
            "                                        )",
            "               Plaintiff,               )",
            "                                        )",
            "          v.                            )",
            "                                        )",
            f"{defendant},",
            "                                        )",
            "               Defendant.               )",
            "________________________________________)",
        ]

        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run = left_para.add_run("\n".join(parties_lines))
        left_run.font.name = FORMAT_SPECS["font_name"]
        left_run.font.size = Pt(FORMAT_SPECS["font_size"])
        left_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Build case info block
        case_number = case_info.get("case_number", "")
        case_display = case_number if case_number else "[Case No. TBD]"
        judge_name = case_info.get("judge_name", "")

        case_lines = ["", "", "", "", f"Case No. {case_display}", "", "", ""]
        if judge_name:
            case_lines.append(f"Judge: {judge_name}")

        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        right_run = right_para.add_run("\n".join(case_lines))
        right_run.font.name = FORMAT_SPECS["font_name"]
        right_run.font.size = Pt(FORMAT_SPECS["font_size"])
        right_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        doc.add_paragraph()

    def _add_document_title(self, doc: Document, doc_type: str, custom_title: str = ""):
        """Add the document title."""
        doc_config = DOCUMENT_TYPES.get(doc_type, DOCUMENT_TYPES.get("motion_to_dismiss"))

        if custom_title:
            title = custom_title.upper()
        elif doc_config:
            title = doc_config["title"]
        else:
            title = "MOTION"

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.name = FORMAT_SPECS["font_name"]
        title_run.font.size = Pt(FORMAT_SPECS["font_size"])

        doc.add_paragraph()

    def _add_body_content(self, doc: Document, sections: Dict[str, str], case_info: Dict[str, str]):
        """Add document body content with proper formatting."""
        # Introduction
        if sections.get("introduction"):
            party_name = case_info.get("plaintiff", "Plaintiff")
            intro_text = f"{party_name}, by and through undersigned counsel, respectfully submits this filing as follows:"

            intro_para = doc.add_paragraph()
            intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            intro_run = intro_para.add_run(intro_text)
            intro_run.font.name = FORMAT_SPECS["font_name"]
            intro_run.font.size = Pt(FORMAT_SPECS["font_size"])
            intro_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

            doc.add_paragraph()
            self._add_content_paragraph(doc, sections["introduction"])

        # Facts
        if sections.get("facts"):
            self._add_section_heading(doc, "FACTUAL BACKGROUND")
            self._add_content_paragraph(doc, sections["facts"])

        # Legal Standard
        if sections.get("legal_standard"):
            self._add_section_heading(doc, "LEGAL STANDARD")
            self._add_content_paragraph(doc, sections["legal_standard"])

        # Argument
        if sections.get("argument"):
            self._add_section_heading(doc, "ARGUMENT")
            self._add_content_paragraph(doc, sections["argument"])

        # Conclusion
        if sections.get("conclusion"):
            self._add_section_heading(doc, "CONCLUSION")
            self._add_content_paragraph(doc, sections["conclusion"])

    def _add_section_heading(self, doc: Document, heading: str):
        """Add a centered, bold section heading."""
        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading_para.add_run(heading)
        heading_run.bold = True
        heading_run.font.name = FORMAT_SPECS["font_name"]
        heading_run.font.size = Pt(FORMAT_SPECS["font_size"])

    def _add_content_paragraph(self, doc: Document, content: str):
        """Add a justified, double-spaced content paragraph."""
        # Split by double newlines to preserve paragraph breaks
        paragraphs = content.split("\n\n")

        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = para.add_run(para_text.strip())
                run.font.name = FORMAT_SPECS["font_name"]
                run.font.size = Pt(FORMAT_SPECS["font_size"])
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                para.paragraph_format.first_line_indent = Inches(0.5)

    def _add_signature_block(self, doc: Document, attorney_info: Dict[str, str], case_info: Dict[str, str]):
        """Add signature block per LCvR 5.1(d)."""
        doc.add_paragraph()

        date_str = datetime.now().strftime("%B %d, %Y")

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]

        # Left: Date
        date_para = left_cell.paragraphs[0]
        date_para.add_run(f"Dated: {date_str}")
        date_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Right: Signature block
        sig_lines = [
            "Respectfully submitted,",
            "",
            f"/s/ {attorney_info.get('name', 'Attorney Name')}",
            attorney_info.get('name', 'Attorney Name'),
        ]

        if attorney_info.get('firm'):
            sig_lines.append(attorney_info['firm'])
        if attorney_info.get('address'):
            sig_lines.append(attorney_info['address'])
        if attorney_info.get('city_state_zip'):
            sig_lines.append(attorney_info['city_state_zip'])
        if attorney_info.get('phone'):
            sig_lines.append(f"Tel: {attorney_info['phone']}")
        if attorney_info.get('email'):
            sig_lines.append(f"Email: {attorney_info['email']}")
        if attorney_info.get('dc_bar_number'):
            sig_lines.append(f"DC Bar No. {attorney_info['dc_bar_number']}")

        party_represented = case_info.get("party_represented", "Plaintiff")
        sig_lines.extend(["", f"Counsel for {party_represented}"])

        sig_para = right_cell.paragraphs[0]
        sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for line in sig_lines:
            run = sig_para.add_run(line + "\n")
            run.font.name = FORMAT_SPECS["font_name"]
            run.font.size = Pt(FORMAT_SPECS["font_size"])
        sig_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _add_certificate_of_service(self, doc: Document, attorney_info: Dict[str, str], case_info: Dict[str, str]):
        """Add certificate of service per LCvR 5.3."""
        doc.add_page_break()

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("CERTIFICATE OF SERVICE")
        title_run.bold = True
        title_run.font.name = FORMAT_SPECS["font_name"]
        title_run.font.size = Pt(FORMAT_SPECS["font_size"])

        doc.add_paragraph()

        # Certificate text
        date_str = datetime.now().strftime("%B %d, %Y")
        cert_text = f"I hereby certify that on {date_str}, a copy of the foregoing document was served via the Court's CM/ECF system on all counsel of record."

        cert_para = doc.add_paragraph()
        cert_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cert_run = cert_para.add_run(cert_text)
        cert_run.font.name = FORMAT_SPECS["font_name"]
        cert_run.font.size = Pt(FORMAT_SPECS["font_size"])
        cert_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        cert_para.paragraph_format.first_line_indent = Inches(0.5)

        doc.add_paragraph()
        doc.add_paragraph()

        # Signature
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_lines = [
            f"/s/ {attorney_info.get('name', 'Attorney Name')}",
            attorney_info.get('name', 'Attorney Name'),
        ]
        for line in sig_lines:
            run = sig_para.add_run(line + "\n")
            run.font.name = FORMAT_SPECS["font_name"]
            run.font.size = Pt(FORMAT_SPECS["font_size"])

    def _add_page_numbers(self, doc: Document):
        """Add page numbers to the footer (bottom center)."""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        run.font.name = FORMAT_SPECS["font_name"]
        run.font.size = Pt(FORMAT_SPECS["font_size"])


# Convenience function
def process_and_reformat(
    file_bytes: bytes,
    doc_type: str,
    case_info: Dict[str, str],
    attorney_info: Optional[Dict[str, str]] = None,
    include_certificate: bool = True
) -> tuple[Dict[str, Any], bytes]:
    """
    Process an uploaded document and reformat it to DC standards.

    Returns:
        Tuple of (extracted_content, reformatted_docx_bytes)
    """
    processor = DocumentProcessor()
    content = processor.extract_from_docx(file_bytes)

    reformatted = processor.reformat_to_dc_standards(
        content=content,
        doc_type=doc_type,
        case_info=case_info,
        attorney_info=attorney_info,
        include_certificate=include_certificate
    )

    return content, reformatted
