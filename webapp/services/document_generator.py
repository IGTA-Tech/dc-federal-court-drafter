"""
Document Generator Service

Generates DOCX and PDF documents with proper DC Federal District Court formatting.
Implements LCvR 5.1 and LCvR 7 requirements.
"""
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))
from config import FORMAT_SPECS, DC_COURT_NAME, DOCUMENT_TYPES


class DocumentGenerator:
    """Generates court documents in DOCX and PDF formats."""

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

    def generate(self, data: dict, format: str = "docx") -> tuple[str, bytes]:
        """
        Generate a court document.

        Args:
            data: Document data including case info, content, etc.
            format: Output format ("docx" or "pdf")

        Returns:
            Tuple of (filename, file_bytes)
        """
        if format.lower() == "docx":
            return self._generate_docx(data)
        elif format.lower() == "pdf":
            return self._generate_pdf(data)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _generate_docx(self, data: dict) -> tuple[str, bytes]:
        """Generate a DOCX document with proper formatting."""
        doc = Document()

        # Set up document formatting
        self._setup_document_format(doc)

        # Add caption
        self._add_caption(doc, data)

        # Add document title
        self._add_document_title(doc, data)

        # Add body content
        self._add_body_content(doc, data)

        # Add signature block
        self._add_signature_block(doc, data)

        # Add certificate of service
        if data.get("include_certificate_of_service", True):
            self._add_certificate_of_service(doc, data)

        # Add page numbers
        self._add_page_numbers(doc)

        # Save to bytes
        filename = self._generate_filename(data, "docx")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return filename, buffer.getvalue()

    def _setup_document_format(self, doc: Document):
        """Set up document formatting per LCvR 7(o)(1)."""
        # Get default section
        section = doc.sections[0]

        # Set page size to Letter
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        # Set margins to 1 inch
        section.top_margin = Inches(FORMAT_SPECS["margin_inches"])
        section.bottom_margin = Inches(FORMAT_SPECS["margin_inches"])
        section.left_margin = Inches(FORMAT_SPECS["margin_inches"])
        section.right_margin = Inches(FORMAT_SPECS["margin_inches"])

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = FORMAT_SPECS["font_name"]
        font.size = Pt(FORMAT_SPECS["font_size"])

        # Set line spacing to double
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    def _add_caption(self, doc: Document, data: dict):
        """Add the caption block per LCvR 5.1(b)(c)."""
        # Court name - centered, bold
        court_para = doc.add_paragraph()
        court_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        court_run = court_para.add_run(DC_COURT_NAME)
        court_run.bold = True
        court_run.font.name = FORMAT_SPECS["font_name"]
        court_run.font.size = Pt(FORMAT_SPECS["font_size"])

        # Add spacing
        doc.add_paragraph()

        # Create the caption box with parties and case number
        # Using a table for layout
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False

        # Set column widths
        table.columns[0].width = Inches(3.5)
        table.columns[1].width = Inches(2.5)

        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]

        # Build left side content (parties)
        parties_text = self._build_parties_block(data)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run = left_para.add_run(parties_text)
        left_run.font.name = FORMAT_SPECS["font_name"]
        left_run.font.size = Pt(FORMAT_SPECS["font_size"])
        # Single spacing for caption
        left_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Build right side content (case number, judge)
        right_text = self._build_case_info_block(data)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        right_run = right_para.add_run(right_text)
        right_run.font.name = FORMAT_SPECS["font_name"]
        right_run.font.size = Pt(FORMAT_SPECS["font_size"])
        right_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Add spacing after caption
        doc.add_paragraph()

    def _build_parties_block(self, data: dict) -> str:
        """Build the parties portion of the caption."""
        plaintiff = data.get("plaintiff", "PLAINTIFF NAME")
        defendant = data.get("defendant", "DEFENDANT NAME")

        lines = [
            "________________________________________",
            "                                        )",
            f"{plaintiff.upper()},",
            "                                        )",
            "               Plaintiff,               )",
            "                                        )",
            "          v.                            )",
            "                                        )",
            f"{defendant.upper()},",
            "                                        )",
            "               Defendant.               )",
            "________________________________________)",
        ]
        return "\n".join(lines)

    def _build_case_info_block(self, data: dict) -> str:
        """Build the case information portion of the caption."""
        case_number = data.get("case_number", "")
        judge_name = data.get("judge_name", "")

        # Use placeholder if case number not provided
        case_display = case_number if case_number else "[Case No. TBD]"

        lines = [
            "",
            "",
            "",
            "",
            f"Case No. {case_display}",
            "",
            "",
            "",
        ]
        if judge_name:
            lines.append(f"Judge: {judge_name}")

        return "\n".join(lines)

    def _add_document_title(self, doc: Document, data: dict):
        """Add the document title."""
        doc_type = data.get("document_type", "motion_to_dismiss")
        doc_config = DOCUMENT_TYPES.get(doc_type, DOCUMENT_TYPES["motion_to_dismiss"])

        title = doc_config["title"]

        # Handle placeholders for opposition/reply
        if "{motion_type}" in title:
            motion_type = data.get("motion_type", "DEFENDANT'S MOTION TO DISMISS")
            title = title.replace("{motion_type}", motion_type)

        # Add custom title if provided
        custom_title = data.get("custom_title", "")
        if custom_title:
            title = custom_title.upper()

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.name = FORMAT_SPECS["font_name"]
        title_run.font.size = Pt(FORMAT_SPECS["font_size"])

        # Add spacing
        doc.add_paragraph()

    def _add_body_content(self, doc: Document, data: dict):
        """Add the document body content."""
        sections = data.get("sections", {})

        # Introduction
        if sections.get("introduction"):
            self._add_section(doc, None, sections["introduction"], is_intro=True, data=data)

        # Background/Facts
        if sections.get("facts"):
            self._add_section(doc, "FACTUAL BACKGROUND", sections["facts"])

        # Legal Standard
        if sections.get("legal_standard"):
            self._add_section(doc, "LEGAL STANDARD", sections["legal_standard"])

        # Argument
        if sections.get("argument"):
            self._add_section(doc, "ARGUMENT", sections["argument"])

        # Additional arguments (A, B, C, etc.)
        for i, arg in enumerate(sections.get("additional_arguments", [])):
            letter = chr(65 + i)  # A, B, C, ...
            heading = arg.get("heading", f"Argument {letter}")
            content = arg.get("content", "")
            self._add_subsection(doc, f"{letter}. {heading}", content)

        # Conclusion
        if sections.get("conclusion"):
            self._add_section(doc, "CONCLUSION", sections["conclusion"])

    def _add_section(self, doc: Document, heading: Optional[str], content: str, is_intro: bool = False, data: dict = None):
        """Add a section with heading and content."""
        if heading:
            heading_para = doc.add_paragraph()
            heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_run = heading_para.add_run(heading)
            heading_run.bold = True
            heading_run.font.name = FORMAT_SPECS["font_name"]
            heading_run.font.size = Pt(FORMAT_SPECS["font_size"])

        # Add introductory language for motions
        if is_intro and data:
            party_name = data.get("party_name", data.get("plaintiff", "Plaintiff"))
            intro_text = f"{party_name}, by and through undersigned counsel, respectfully moves this Court as follows:"
            intro_para = doc.add_paragraph()
            intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            intro_run = intro_para.add_run(intro_text)
            intro_run.font.name = FORMAT_SPECS["font_name"]
            intro_run.font.size = Pt(FORMAT_SPECS["font_size"])
            intro_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            doc.add_paragraph()

        # Add content paragraphs
        content_para = doc.add_paragraph()
        content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        content_run = content_para.add_run(content)
        content_run.font.name = FORMAT_SPECS["font_name"]
        content_run.font.size = Pt(FORMAT_SPECS["font_size"])
        content_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    def _add_subsection(self, doc: Document, heading: str, content: str):
        """Add a subsection (A., B., etc.)."""
        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_run = heading_para.add_run(heading)
        heading_run.bold = True
        heading_run.font.name = FORMAT_SPECS["font_name"]
        heading_run.font.size = Pt(FORMAT_SPECS["font_size"])

        content_para = doc.add_paragraph()
        content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        content_run = content_para.add_run(content)
        content_run.font.name = FORMAT_SPECS["font_name"]
        content_run.font.size = Pt(FORMAT_SPECS["font_size"])
        content_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    def _add_signature_block(self, doc: Document, data: dict):
        """Add signature block per LCvR 5.1(d)."""
        attorney = data.get("attorney", {})

        doc.add_paragraph()  # Spacing

        # Date line
        date_str = data.get("date", datetime.now().strftime("%B %d, %Y"))

        # Create two-column layout for date and signature
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]

        # Left: Date
        date_para = left_cell.paragraphs[0]
        date_para.add_run(f"Dated: {date_str}")
        date_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Right: Signature block
        sig_lines = [
            "Respectfully submitted,",
            "",
            f"/s/ {attorney.get('name', 'Attorney Name')}",
            attorney.get('name', 'Attorney Name'),
        ]

        if attorney.get('firm'):
            sig_lines.append(attorney['firm'])

        if attorney.get('address'):
            sig_lines.append(attorney['address'])

        if attorney.get('city_state_zip'):
            sig_lines.append(attorney['city_state_zip'])

        if attorney.get('phone'):
            sig_lines.append(f"Tel: {attorney['phone']}")

        if attorney.get('email'):
            sig_lines.append(f"Email: {attorney['email']}")

        if attorney.get('dc_bar_number'):
            sig_lines.append(f"DC Bar No. {attorney['dc_bar_number']}")

        party_represented = data.get("party_represented", "Plaintiff")
        sig_lines.append("")
        sig_lines.append(f"Counsel for {party_represented}")

        sig_para = right_cell.paragraphs[0]
        sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for line in sig_lines:
            run = sig_para.add_run(line + "\n")
            run.font.name = FORMAT_SPECS["font_name"]
            run.font.size = Pt(FORMAT_SPECS["font_size"])
        sig_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _add_certificate_of_service(self, doc: Document, data: dict):
        """Add certificate of service per LCvR 5.3."""
        doc.add_page_break()

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("CERTIFICATE OF SERVICE")
        title_run.bold = True
        title_run.font.name = FORMAT_SPECS["font_name"]
        title_run.font.size = Pt(FORMAT_SPECS["font_size"])

        doc.add_paragraph()

        # Certificate text
        date_str = data.get("date", datetime.now().strftime("%B %d, %Y"))
        doc_title = data.get("custom_title", "foregoing document")

        cert_text = f"I hereby certify that on {date_str}, a copy of the {doc_title} was served via the Court's CM/ECF system on all counsel of record."

        cert_para = doc.add_paragraph()
        cert_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cert_run = cert_para.add_run(cert_text)
        cert_run.font.name = FORMAT_SPECS["font_name"]
        cert_run.font.size = Pt(FORMAT_SPECS["font_size"])
        cert_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        cert_para.paragraph_format.first_line_indent = Inches(0.5)

        doc.add_paragraph()
        doc.add_paragraph()

        # Signature
        attorney = data.get("attorney", {})
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_lines = [
            f"/s/ {attorney.get('name', 'Attorney Name')}",
            attorney.get('name', 'Attorney Name'),
        ]
        for line in sig_lines:
            run = sig_para.add_run(line + "\n")
            run.font.name = FORMAT_SPECS["font_name"]
            run.font.size = Pt(FORMAT_SPECS["font_size"])

    def _add_page_numbers(self, doc: Document):
        """Add page numbers to the footer (bottom center)."""
        section = doc.sections[0]

        footer = section.footer
        footer.is_linked_to_previous = False

        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        run.font.name = FORMAT_SPECS["font_name"]
        run.font.size = Pt(FORMAT_SPECS["font_size"])

    def _generate_filename(self, data: dict, extension: str) -> str:
        """Generate a filename for the document."""
        case_number = data.get("case_number", "")
        # Use 'draft' if no case number provided
        case_clean = re.sub(r'[^\w\-]', '_', case_number) if case_number else "draft"

        doc_type = data.get("document_type", "document")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        return f"{case_clean}_{doc_type}_{timestamp}.{extension}"

    def _generate_pdf(self, data: dict) -> tuple[str, bytes]:
        """Generate a text-searchable PDF document."""
        buffer = BytesIO()

        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )

        # Create styles
        styles = getSampleStyleSheet()

        # Court header style
        court_style = ParagraphStyle(
            'CourtHeader',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=12,
            alignment=TA_CENTER,
            spaceAfter=12,
            leading=14
        )

        # Title style
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=12,
            alignment=TA_CENTER,
            spaceAfter=24,
            spaceBefore=12
        )

        # Heading style
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=12,
            alignment=TA_CENTER,
            spaceAfter=12,
            spaceBefore=24
        )

        # Body style (double-spaced)
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=12,
            alignment=TA_JUSTIFY,
            leading=24,  # Double spacing
            firstLineIndent=36,
            spaceAfter=0
        )

        # Caption style (single-spaced)
        caption_style = ParagraphStyle(
            'Caption',
            parent=styles['Normal'],
            fontName='Courier',
            fontSize=10,
            alignment=TA_LEFT,
            leading=12,
            spaceAfter=12
        )

        # Signature style
        sig_style = ParagraphStyle(
            'Signature',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=12,
            alignment=TA_LEFT,
            leading=14,
            leftIndent=216  # 3 inches from left
        )

        story = []

        # Court name
        court_text = DC_COURT_NAME.replace("\n", "<br/>")
        story.append(Paragraph(f"<b>{court_text}</b>", court_style))
        story.append(Spacer(1, 12))

        # Caption
        caption_text = self._build_pdf_caption(data)
        story.append(Paragraph(caption_text.replace("\n", "<br/>"), caption_style))
        story.append(Spacer(1, 12))

        # Document title
        doc_type = data.get("document_type", "motion_to_dismiss")
        doc_config = DOCUMENT_TYPES.get(doc_type, DOCUMENT_TYPES["motion_to_dismiss"])
        title = doc_config["title"]
        if "{motion_type}" in title:
            motion_type = data.get("motion_type", "DEFENDANT'S MOTION TO DISMISS")
            title = title.replace("{motion_type}", motion_type)
        custom_title = data.get("custom_title", "")
        if custom_title:
            title = custom_title.upper()

        story.append(Paragraph(f"<b>{title}</b>", title_style))

        # Sections
        sections = data.get("sections", {})

        # Introduction
        if sections.get("introduction"):
            party_name = data.get("party_name", data.get("plaintiff", "Plaintiff"))
            intro = f"{party_name}, by and through undersigned counsel, respectfully moves this Court as follows:"
            story.append(Paragraph(intro, body_style))
            story.append(Spacer(1, 12))
            story.append(Paragraph(sections["introduction"], body_style))

        # Facts
        if sections.get("facts"):
            story.append(Paragraph("<b>FACTUAL BACKGROUND</b>", heading_style))
            story.append(Paragraph(sections["facts"], body_style))

        # Legal Standard
        if sections.get("legal_standard"):
            story.append(Paragraph("<b>LEGAL STANDARD</b>", heading_style))
            story.append(Paragraph(sections["legal_standard"], body_style))

        # Argument
        if sections.get("argument"):
            story.append(Paragraph("<b>ARGUMENT</b>", heading_style))
            story.append(Paragraph(sections["argument"], body_style))

        # Conclusion
        if sections.get("conclusion"):
            story.append(Paragraph("<b>CONCLUSION</b>", heading_style))
            story.append(Paragraph(sections["conclusion"], body_style))

        # Signature block
        story.append(Spacer(1, 36))
        sig_block = self._build_pdf_signature(data)
        story.append(Paragraph(sig_block.replace("\n", "<br/>"), sig_style))

        # Certificate of service
        if data.get("include_certificate_of_service", True):
            story.append(PageBreak())
            story.append(Paragraph("<b>CERTIFICATE OF SERVICE</b>", title_style))
            date_str = data.get("date", datetime.now().strftime("%B %d, %Y"))
            cert_text = f"I hereby certify that on {date_str}, a copy of the foregoing document was served via the Court's CM/ECF system on all counsel of record."
            story.append(Paragraph(cert_text, body_style))
            story.append(Spacer(1, 36))
            attorney = data.get("attorney", {})
            sig = f"/s/ {attorney.get('name', 'Attorney Name')}<br/>{attorney.get('name', 'Attorney Name')}"
            story.append(Paragraph(sig, sig_style))

        # Build PDF
        doc.build(story, onFirstPage=self._add_pdf_page_number, onLaterPages=self._add_pdf_page_number)

        buffer.seek(0)
        filename = self._generate_filename(data, "pdf")

        return filename, buffer.getvalue()

    def _build_pdf_caption(self, data: dict) -> str:
        """Build caption text for PDF."""
        plaintiff = data.get("plaintiff", "PLAINTIFF NAME").upper()
        defendant = data.get("defendant", "DEFENDANT NAME").upper()
        case_number = data.get("case_number", "")
        # Use placeholder if case number not provided
        case_display = case_number if case_number else "[Case No. TBD]"
        judge_name = data.get("judge_name", "")

        lines = [
            "________________________________________",
            "                                        )",
            f"{plaintiff},                           )",
            "                                        )",
            "               Plaintiff,               )    Case No. " + case_display,
            "                                        )",
            "          v.                            )",
            "                                        )",
            f"{defendant},                           )",
            "                                        )",
            "               Defendant.               )",
            "________________________________________)   " + (f"Judge: {judge_name}" if judge_name else ""),
        ]
        return "\n".join(lines)

    def _build_pdf_signature(self, data: dict) -> str:
        """Build signature block text for PDF."""
        attorney = data.get("attorney", {})
        date_str = data.get("date", datetime.now().strftime("%B %d, %Y"))

        lines = [
            f"Dated: {date_str}",
            "",
            "Respectfully submitted,",
            "",
            f"/s/ {attorney.get('name', 'Attorney Name')}",
            attorney.get('name', 'Attorney Name'),
        ]

        if attorney.get('firm'):
            lines.append(attorney['firm'])
        if attorney.get('address'):
            lines.append(attorney['address'])
        if attorney.get('city_state_zip'):
            lines.append(attorney['city_state_zip'])
        if attorney.get('phone'):
            lines.append(f"Tel: {attorney['phone']}")
        if attorney.get('email'):
            lines.append(f"Email: {attorney['email']}")
        if attorney.get('dc_bar_number'):
            lines.append(f"DC Bar No. {attorney['dc_bar_number']}")

        party_represented = data.get("party_represented", "Plaintiff")
        lines.append("")
        lines.append(f"Counsel for {party_represented}")

        return "\n".join(lines)

    def _add_pdf_page_number(self, canvas, doc):
        """Add page numbers to PDF."""
        canvas.saveState()
        canvas.setFont('Times-Roman', 12)
        page_num = canvas.getPageNumber()
        text = str(page_num)
        canvas.drawCentredString(letter[0] / 2, 0.5 * inch, text)
        canvas.restoreState()


# Convenience function
def generate_document(data: dict, format: str, output_dir: str) -> tuple[str, bytes]:
    """Generate a court document."""
    generator = DocumentGenerator(output_dir)
    return generator.generate(data, format)
